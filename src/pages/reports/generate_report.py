import dash
from dash import html, callback, Output, Input, State, ctx, dcc
from utils import data_loader
import pandas as pd
from urllib.parse import parse_qs
import numpy as np
import dash_bootstrap_components as dbc
import io
from docx import Document
from bs4 import BeautifulSoup

dash.register_page(__name__, path='/reports/generate-report')

def layout(report_type=None, institution_type=None, **kwargs):
	layout = html.Div([
		dcc.Location(id='generate-report-url'),
		html.Div(
			[
				dbc.Collapse(
					dbc.Card(
						dbc.CardBody([
							dbc.Button("Return to Selection", id="generate-report-previous-btn", color="light",
									   className="w-100"),
							dbc.Button("Download Report", id="download-btn", color="success", className="w-100"),
						], className="d-grid gap-3"),
						className="mb-2 shadow"
					),
					id="report-fab-collapse",
					is_open=False
				),
				dbc.Button(
					"⋮",
					id="report-fab-toggle",
					color="primary",
					className="rounded-circle fw-bold",
					style={
						"width": "56px",
						"height": "56px",
						"fontSize": "1.5rem",
						"boxShadow": "0 4px 12px rgba(0,0,0,0.3)"
					}
				),
			],
			className="position-fixed d-flex flex-column align-items-end",
			style={"bottom": "2rem", "right": "2rem", "zIndex": 1050}
		),
		dbc.Row([
			dbc.Col(
				html.Div(id='report-sidebar'),
				style={
					'height': 'calc(100vh - 70px)',
					'overflowY': 'auto',
					'position': 'sticky',
					'top': '120px'
				},
				md=3,
				className="sidebar-btn-group mb-3 bg-light border-0"
			),
			dbc.Col(html.Div(id='report-content'), className="ms-22 pt-80"),
		]),
		dcc.Download(id="download-store"),
	], className="container")
	return layout

@callback(
	Output("report-fab-collapse", "is_open"),
	Input("report-fab-toggle", "n_clicks"),
	State("report-fab-collapse", "is_open"),
	prevent_initial_call=True
)
def toggle_report_fab(n, is_open):
	return not is_open if n else is_open

@callback(
	Output("generate-report-url", "href"),
	Input("user-selection-completed-store", "data"),
	Input("generate-report-url", "search"),
	Input("generate-report-previous-btn", "n_clicks"),
	State("report-type-store", "data"),
	State("institution-type-store", "data"),
	State("generate-report-url", "pathname"),
)
def update_url(user_selection_completed, url_search, back, report_type, institution_type, url_pathname):
	query = parse_qs(url_search.lstrip('?'))
	start_page = len(query) == 0 and "generate-report" in url_pathname
	report_page = query.get('report_type', [None])[0]

	if start_page and user_selection_completed and not report_page:
		return f"/reports/generate-report?report-type={report_type}&institution-type={institution_type}"
	if back:
		return f"/reports/customise-report?report_type={report_type}&institution-type={institution_type}&review=summary"
	return dash.no_update

@callback(
	Output("download-store", "data"),
	Input("download-btn", "n_clicks"),
	State("report-content", "children"),
	prevent_initial_call=True,
)
def download_report(n_clicks, report_content):
	# html_string = _dash_html_to_string(report_content)
	# doc = _html_to_docx(html_string)
	# buffer = io.BytesIO()
	# doc.save(buffer)
	# buffer.seek(0)
	# return dict(
	# 	content=buffer.read(),
	# 	filename="CFRF_Report.docx",
	# 	type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
	# )

	html_string = _dash_html_to_string(report_content)
	# Convert HTML to plain text for RTF (or use a library for rich formatting)
	soup = BeautifulSoup(html_string, "html.parser")
	plain_text = soup.get_text()
	rtf_content = r"{\rtf1\ansi " + plain_text.replace('\n', r'\par ') + "}"
	buffer = io.BytesIO()
	buffer.write(rtf_content.encode('utf-8'))
	buffer.seek(0)
	return dict(
		content=buffer.read(),
		filename="CFRF_Report.rtf",
		type="application/rtf"
	)

@callback(
	Output("report-sidebar", "children"),
	Output("report-content", "children"),
	Input("generate-report-url", "search"),
	Input("all-user-selection-store", "data"),
	State("output-structure-mapping-store", "data"),
	State("exposure-sector-product-mapping-store", "data"),
	State("scenario-mapping-store", "data"),
	State("report-type-store", "data"),
	State("generate-report-url", "pathname"),
	prevent_initial_call=True
)
def generate_all_reports(url_search, all_stored_data, output_structure_mapping_dict,
						 exposure_sector_product_mapping_dict, scenario_mapping_dict, report_type, url_pathname):
	query = parse_qs(url_search.lstrip('?'))
	start_page = len(query) == 0 and "generate-report" in url_pathname

	if start_page:
		raise dash.exceptions.PreventUpdate

	# Initialise all dataframes
	all_user_selection_df = pd.DataFrame()
	for k, v in all_stored_data.items():
		user_selection_df = pd.DataFrame(all_stored_data[k])
		all_user_selection_df = pd.concat([all_user_selection_df, user_selection_df])

	scenario_mapping_df = pd.DataFrame(scenario_mapping_dict)
	exposure_sector_product_mapping_df = pd.DataFrame(exposure_sector_product_mapping_dict)
	output_structure_mapping_df = pd.DataFrame(output_structure_mapping_dict)

	# Extract scenario list if any
	scenario_list = list(all_user_selection_df[all_user_selection_df['exposure'] == 'Scenario']['label'].unique())

	# Prepare output structure mapping
	sorted_output_structure_mapping_df = prepare_output_structure_mapping(output_structure_mapping_df, report_type)

	# Get user selection with sector and product yml file path
	user_selection_with_yml_file_path_df = get_sector_and_product_yml_file_path(
		all_user_selection_df, sorted_output_structure_mapping_df, exposure_sector_product_mapping_df
	)

	# Get sort order
	user_selection_with_yml_file_path_df = create_sort_order(user_selection_with_yml_file_path_df)

	# Output structure
	output_structure_layout = []
	for section in sorted_output_structure_mapping_df['output_structure'].unique():
		section_layout = []
		section_mapping_df = sorted_output_structure_mapping_df[
			sorted_output_structure_mapping_df['output_structure'] == section
		]
		for sub_section in section_mapping_df['sub_section'].unique():
			sub_section_mapping_df = section_mapping_df[section_mapping_df['sub_section'] == sub_section]
			if sub_section == 'summary_input_table_flag':
				sub_section_layout = get_summary_input_table_layout(
					all_user_selection_df, user_selection_with_yml_file_path_df
				)
			elif sub_section == 'scenario_content_id':
				sub_section_layout = get_scenario_layout(sub_section_mapping_df, scenario_mapping_df, scenario_list)
			elif sub_section == 'sector_description_content_id':
				filtered_user_selection_with_yml_file_path_df = _filter_user_selection_by_section(
					user_selection_with_yml_file_path_df, section, sub_section
				)
				sub_section_layout = get_sector_description_layout(filtered_user_selection_with_yml_file_path_df)
			elif sub_section == 'sector_scenario_content_id':
				filtered_user_selection_with_yml_file_path_df = _filter_user_selection_by_section(
					user_selection_with_yml_file_path_df, section, sub_section
				)
				sub_section_layout = get_sector_scenario_layout(
					filtered_user_selection_with_yml_file_path_df, scenario_mapping_df, scenario_list
				)
			elif sub_section == 'product_content_id':
				filtered_user_selection_with_yml_file_path_df = _filter_user_selection_by_section(
					user_selection_with_yml_file_path_df, section, sub_section
				)
				sub_section_layout = get_product_text_layout(filtered_user_selection_with_yml_file_path_df)
			else:
				sub_section_layout = []
			section_layout.append(sub_section_layout)
		section_layout = html.Div([html.H1(section)] + section_layout)
		output_structure_layout.append(section_layout)
	output_structure_layout = remove_parents_with_empty_children(output_structure_layout)
	output_structure_layout = clean_up_sector_overview_and_detail(output_structure_layout)
	output_structure_layout, sidebar_layout = create_sidebar_layout(output_structure_layout)
	return html.Div(sidebar_layout), html.Div(output_structure_layout)

def prepare_output_structure_mapping(output_structure_mapping_df, report_type):
	exploded_output_structure_mapping_df = output_structure_mapping_df.copy()
	exploded_output_structure_mapping_df['materiality'] = [
		['Low', 'Medium', 'High'] if x == 'All' else [x] for x in exploded_output_structure_mapping_df['materiality']
	]
	exploded_output_structure_mapping_df = exploded_output_structure_mapping_df.explode(['materiality'])
	exploded_output_structure_mapping_df['materiality'] = pd.Categorical(
		exploded_output_structure_mapping_df['materiality'],
		categories=['Low', 'Medium', 'High'],
		ordered=True
	)
	output_structure_order_df = exploded_output_structure_mapping_df[['report_type', 'output_structure', 'materiality']].drop_duplicates()

	melted_output_structure_mapping_df = pd.melt(
		exploded_output_structure_mapping_df,
		id_vars=['report_type', 'output_structure', 'materiality'],
		var_name='sub_section', value_name='content_id'
	)
	melted_output_structure_mapping_df = melted_output_structure_mapping_df[
		~(melted_output_structure_mapping_df['content_id'].isnull())
		& (melted_output_structure_mapping_df['report_type'] == report_type)
	]
	sorted_output_structure_mapping_df = pd.merge(
		output_structure_order_df, melted_output_structure_mapping_df,
		on=['report_type', 'output_structure', 'materiality'], how='left'
	)
	return sorted_output_structure_mapping_df

def get_sector_and_product_yml_file_path(all_user_selection_df, output_structure_mapping_df, exposure_sector_product_mapping_df):
	user_selection_df = all_user_selection_df[all_user_selection_df['exposure'] != 'Scenario'].copy()
	user_selection_df = user_selection_df.rename(columns={
		'report': 'report_type',
		'label': 'materiality',
		'ptype': 'type',
	})
	user_selection_df = user_selection_df.drop(columns=['id', 'value'])

	# Merge in exposure-sector-product mapping to get yml file
	sovereign_user_selection_df = user_selection_df[user_selection_df['exposure'] == 'Sovereign']
	sovereign_user_selection_with_mapping_df = pd.merge(
		sovereign_user_selection_df, exposure_sector_product_mapping_df.drop(columns=['institution', 'type']),
		on=['exposure', 'sector'], how='left'
	)
	other_user_selection_df = user_selection_df[user_selection_df['exposure'] != 'Sovereign']
	other_user_selection_with_mapping_df = pd.merge(
		other_user_selection_df, exposure_sector_product_mapping_df, on=['institution', 'exposure', 'sector', 'type'], how='left'
	)
	user_selection_with_mapping_df = pd.concat([sovereign_user_selection_with_mapping_df, other_user_selection_with_mapping_df])
	user_selection_with_mapping_df = user_selection_with_mapping_df[user_selection_with_mapping_df['materiality'] != 'N/A']

	# Merge in output structure mapping to get content_id
	user_selection_with_mapping_df = pd.merge(
		user_selection_with_mapping_df, output_structure_mapping_df, on=['report_type', 'materiality'], how='left'
	)
	return user_selection_with_mapping_df

def create_sort_order(user_selection_with_yml_file_path_df):
	output_df = user_selection_with_yml_file_path_df.copy()

	# Get sort order
	output_df['sort_order'] = [
		next(iter(data_loader.load_yml_file('exposure_class', f'{sector_yml_file}.yml')))
		for sector_yml_file in output_df['sector_yml_file']
	]
	output_df['materiality'] = pd.Categorical(
		output_df['materiality'],
		categories=['Low', 'Medium', 'High'],
		ordered=True
	)

	# Get sector name
	output_df['yml_sector'] = [
		next(iter(data_loader.load_yml_file('exposure_class', f'{sector_yml_file}.yml').values()))['name']
		for sector_yml_file in output_df['sector_yml_file']
	]

	# Sort by sort order and materiality
	output_df = output_df.sort_values(by=['sort_order', 'materiality'])
	return output_df

def get_summary_input_table_layout(all_user_selection_df, user_selection_with_yml_file_path_df):
	scenario_user_selection_df = all_user_selection_df[all_user_selection_df['exposure'] == 'Scenario'].copy()
	scenario_user_selection_df = scenario_user_selection_df.rename(columns={
		'report': 'report_type',
		'label': 'materiality',
		'ptype': 'type',
	})
	summary_input_df = pd.concat([user_selection_with_yml_file_path_df, scenario_user_selection_df])
	summary_input_table = data_loader.rename_user_selection_data_columns(summary_input_df).drop_duplicates()
	summary_input_table_layout = html.Div([
		html.H2('Summary of Inputs'),
		data_loader.create_data_table(summary_input_table)
	])
	return summary_input_table_layout

def get_scenario_layout(sub_section_mapping_df, scenario_mapping_df, scenario_list):
	section = sub_section_mapping_df['output_structure'].iloc[0]
	content_id = sub_section_mapping_df['content_id'].iloc[0]
	all_desc = []
	if len(scenario_list) == 0:
		scenario_layout = html.Div(className='d-none')
	else:
		for scenario in scenario_list:
			filtered_scenario_mapping_df = scenario_mapping_df[scenario_mapping_df['scenario_name'] == scenario]
			scenario_yml_file = filtered_scenario_mapping_df['scenario_yml_file'].iloc[0]
			scenario_yml = data_loader.load_yml_file('scenario', f'{scenario_yml_file}.yml')
			content_value = scenario_yml.get(content_id, "")
			desc = html.Div([
				html.H2(scenario) if section == 'Scenario Detail' else html.H3(scenario),
				dcc.Markdown(content_value, link_target="_blank", dangerously_allow_html=True, className='display-12', style={'textTransform': 'none'})
			], className='mb-3')
			all_desc.append(desc)

		if section == 'Executive Summary':
			header_layout = [
				html.H2(f'Summary of Scenario{data_loader.plural_add_s(len(scenario_list) > 0)}'),
				html.P(f'This report considers the following scenario'
					   f'{data_loader.plural_add_s(len(scenario_list) > 0)}:'),
			]
		else:
			header_layout = []

		scenario_layout = html.Div([*header_layout, *all_desc])
	return scenario_layout

def get_sector_scenario_layout(user_selection_with_yml_file_path_df, scenario_mapping_df, scenario_list):
	section = user_selection_with_yml_file_path_df['output_structure'].iloc[0]

	# Split by materiality
	high_materiality_df = user_selection_with_yml_file_path_df[user_selection_with_yml_file_path_df['materiality'] == 'High'].copy()
	high_materiality_df = high_materiality_df[['sector_yml_file', 'content_id', 'materiality']].drop_duplicates()
	low_medium_materiality_df = user_selection_with_yml_file_path_df[user_selection_with_yml_file_path_df['materiality'].isin(['Low', 'Medium'])].copy()
	low_medium_materiality_df = low_medium_materiality_df[['sector_yml_file', 'content_id', 'materiality']].drop_duplicates()

	if len(scenario_list) == 0:
		sector_scenario_layout = html.Div(className='d-none')
	else:
		# High materiality section
		if len(high_materiality_df) == 0:
			high_materiality_layout = html.Div(className='d-none')
			high_materiality_other_desc = []
			high_materiality_sovereign_desc = []
		else:
			high_materiality_sovereign_desc = []
			high_materiality_other_desc = []
			for sector_yml_file in high_materiality_df['sector_yml_file'].unique():
				sector_selection_df = high_materiality_df[high_materiality_df['sector_yml_file'] == sector_yml_file]
				sector_yml_file = sector_selection_df['sector_yml_file'].iloc[0]
				sector_yml = data_loader.load_yml_file('exposure_class', f'{sector_yml_file}.yml')
				sector_yml = next(iter(sector_yml.values()))
				sector_name = sector_yml['name']
				content_id_list = list(sector_selection_df['content_id'].unique())
				content_id_list = sorted(content_id_list, key=lambda x: ['always', 'high_materiality'].index(x))
				scenario_desc = []
				for scenario in scenario_list:
					if 'sovereign' in sector_name.lower():
						header = [html.H6(scenario)]
					else:
						header = [html.H5(scenario)]
					content_desc = []
					for content_id in content_id_list:
						if section == 'Sector Detail':
							if content_id == 'always':
								sub_header = [dcc.Markdown('**Summary**')]
							else:
								sub_header = [dcc.Markdown('**Detail**')]
						else:
							sub_header = []
						content = [dcc.Markdown(
							_filter_yml_by_scenario(sector_yml, scenario, scenario_mapping_df)[content_id],
							link_target="_blank", dangerously_allow_html=True, className='display-12', style={'textTransform': 'none'}
						)]
						content_desc.append(html.Div(sub_header + content))
					scenario_desc.append(html.Div(header + content_desc, className='mb-3'))
				if 'sovereign' in sector_name.lower():
					sector_div = html.Div([html.H5(sector_name), *scenario_desc])
					high_materiality_sovereign_desc.append(sector_div)
				else:
					sector_div = html.Div([html.H4(sector_name), *scenario_desc])
					high_materiality_other_desc.append(sector_div)
			high_materiality_layout = html.Div([
				html.H3('High materiality exposures') if section == 'Executive Summary' else html.Div([], className='d-none'),
				*high_materiality_other_desc,
				html.Div([
					html.H4('Sovereign', className='fw-bold'),
					*high_materiality_sovereign_desc
				]),
			])

		# Low and Medium materiality section
		if len(low_medium_materiality_df) == 0:
			low_medium_materiality_layout = html.Div([], className='d-none')
			low_medium_materiality_other_desc = []
			low_medium_materiality_sovereign_desc = []
		else:
			low_medium_materiality_sovereign_desc = []
			low_medium_materiality_other_desc = []
			if section == 'Executive Summary':
				# Combine materiality
				low_medium_materiality_df = _convert_to_bullet_points(low_medium_materiality_df, 'materiality')

				# Add scenarios
				low_medium_materiality_df['scenario'] = [scenario_list for x in low_medium_materiality_df['sector_yml_file']]
				low_medium_materiality_df = low_medium_materiality_df.explode('scenario')

				# Add sector yml
				low_medium_materiality_df['sector_yml'] = [
					next(iter(data_loader.load_yml_file('exposure_class', f'{sector_yml_file}.yml').values()))
					for sector_yml_file in low_medium_materiality_df['sector_yml_file']
				]
				low_medium_materiality_df['sector'] = [
					sector_yml['name'] for sector_yml in low_medium_materiality_df['sector_yml']
				]

				# Add description
				low_medium_materiality_df['description'] = [
					_filter_yml_by_scenario(sector_yml, scenario, scenario_mapping_df)[content_id]
					for sector_yml, scenario, content_id in low_medium_materiality_df[['sector_yml', 'scenario', 'content_id']].to_numpy()
				]

				# Create datatable
				low_medium_materiality_table = data_loader.rename_user_selection_data_columns(low_medium_materiality_df)
				low_medium_materiality_table = data_loader.create_data_table(low_medium_materiality_table, ['Materiality'], ['Description'])

				low_medium_materiality_layout = html.Div([
					html.H3('Other exposures'),
					html.Div(low_medium_materiality_table)
				])
			elif section == 'Sector Detail':
				low_medium_materiality_sovereign_desc = []
				low_medium_materiality_other_desc = []
				for sector_yml_file in low_medium_materiality_df['sector_yml_file'].unique():
					sector_selection_df = low_medium_materiality_df[low_medium_materiality_df['sector_yml_file'] == sector_yml_file]
					sector_yml_file = sector_selection_df['sector_yml_file'].iloc[0]
					sector_yml = data_loader.load_yml_file('exposure_class', f'{sector_yml_file}.yml')
					sector_yml = next(iter(sector_yml.values()))
					content_id = sector_selection_df['content_id'].iloc[0]
					scenario_desc = [
						html.Div([
							html.H5(scenario),
							dcc.Markdown(
								_filter_yml_by_scenario(sector_yml, scenario, scenario_mapping_df)[content_id],
								link_target="_blank", dangerously_allow_html=True, className='display-12', style={'textTransform': 'none'}
							)
						]) for scenario in scenario_list
					]
					sector_name = sector_yml['name']
					if 'sovereign' in sector_name.lower():
						sector_div = html.Div([html.H5(sector_name), *scenario_desc])
						low_medium_materiality_sovereign_desc.append(sector_div)
					else:
						sector_div = html.Div([html.H4(sector_name), *scenario_desc])
						low_medium_materiality_other_desc.append(sector_div)
				low_medium_materiality_layout = html.Div([
					*low_medium_materiality_other_desc,
					html.Div([
						html.H4('Sovereign', className='fw-bold') if len(low_medium_materiality_sovereign_desc) > 0 else [],
						*low_medium_materiality_sovereign_desc
					]),
				])
			else:
				low_medium_materiality_layout = html.Div(className='d-none')

		if section == 'Executive Summary':
			sector_scenario_layout = html.Div([
				html.H2(f'Summary of Exposure{data_loader.plural_add_s(len(user_selection_with_yml_file_path_df) > 0)}'),
				html.P(f'This report considers the following exposure'
					   f'{data_loader.plural_add_s(len(user_selection_with_yml_file_path_df) > 0)}:'),
				high_materiality_layout,
				low_medium_materiality_layout
			])
		elif section == 'Sector Detail':
			sector_list = []
			other_desc = []
			for desc in high_materiality_other_desc:
				if hasattr(desc, 'children'):
					sector_name = desc.children[0].children
					sector_list.append(sector_name)
					other_desc.append(desc)

			for desc in low_medium_materiality_other_desc:
				if hasattr(desc, 'children'):
					sector_name = desc.children[0].children
					if sector_name not in sector_list:
						sector_list.append(sector_name)
						other_desc.append(desc)

			sector_scenario_layout = html.Div([
				*other_desc,
				html.Div([
					html.H4('Sovereign', className='fw-bold'),
					*high_materiality_sovereign_desc,
					*low_medium_materiality_sovereign_desc,
				])
			])
		else:
			sector_scenario_layout = html.Div(className='d-none')
	return sector_scenario_layout

def get_sector_description_layout(user_selection_with_yml_file_path_df):
	# Split by sovereign and other sectors
	sovereign_sector_selection_df = user_selection_with_yml_file_path_df[
		user_selection_with_yml_file_path_df['exposure'] == 'Sovereign'
	].copy()
	other_sector_selection_df = user_selection_with_yml_file_path_df[
		user_selection_with_yml_file_path_df['exposure'] != 'Sovereign'
	].copy()

	all_sector_desc = []
	for sector_selection_df in [other_sector_selection_df, sovereign_sector_selection_df]:
		if len(sector_selection_df) == 0:
			continue
		else:
			# Get unique sector yml files without materiality
			sector_selection_df = sector_selection_df[['sector_yml_file', 'content_id']].drop_duplicates()

			# Add sector yml
			sector_selection_df['sector_yml_file'] = np.where(
				sector_selection_df['sector_yml_file'].str.contains('sov'),
				'sovereigns/sovereigns', sector_selection_df['sector_yml_file']
			)
			sector_selection_df['sector_yml'] = [
				next(iter(data_loader.load_yml_file('exposure_class', f'{sector_yml_file}.yml').values()))
				for sector_yml_file in sector_selection_df['sector_yml_file']
			]
			sector_selection_df['sector'] = [
				sector_yml['name'] for sector_yml in sector_selection_df['sector_yml']
			]

			# Add description
			sector_selection_df['description'] = [
				sector_yml.get(content_id, "")
				for sector_yml, content_id in sector_selection_df[['sector_yml', 'content_id']].to_numpy()
			]

			# Create layout
			sector_selection_df = sector_selection_df[['sector', 'description']].drop_duplicates()
			sector_desc = [
				html.Div([
					html.H5(sector),
					dcc.Markdown(description, link_target="_blank", dangerously_allow_html=True, className='display-12', style={'textTransform': 'none'})
				]) if description != "" else html.Div(className='d-none')
				for sector, description in sector_selection_df[['sector', 'description']].to_numpy()
			]
			all_sector_desc += sector_desc
	sector_layout = html.Div([*all_sector_desc]) if len(all_sector_desc) > 0 else html.Div(className='d-none')
	return sector_layout

def get_product_text_layout(user_selection_with_yml_file_path_df):
	product_selection_df = user_selection_with_yml_file_path_df.copy()

	if len(product_selection_df) == 0:
		product_layout = html.Div(className='d-none')
	else:
		# Drop materiality
		product_selection_df = product_selection_df.drop(columns='materiality').drop_duplicates()

		# Change ptype to sector if it's Exposure
		product_selection_df['type'] = np.where(
			product_selection_df['type'] == 'Exposure', product_selection_df['sector'], product_selection_df['type']
		)

		# Add product yml
		product_selection_df['product_yml'] = [
			next(iter(data_loader.load_yml_file('product', f'{product_yml_file}.yml').values()))
			for product_yml_file in product_selection_df['product_yml_file']
		]

		# Add product name
		product_selection_df['product'] = [
			product_yml['description'] for product_yml in product_selection_df['product_yml']
		]

		# Add sector yml
		product_selection_df['sector_yml_file'] = np.where(
			product_selection_df['sector_yml_file'].str.contains('sov'),
			'sovereigns/sovereigns', product_selection_df['sector_yml_file']
		)
		product_selection_df['sector_yml'] = [
			next(iter(data_loader.load_yml_file('exposure_class', f'{sector_yml_file}.yml').values()))
			for sector_yml_file in product_selection_df['sector_yml_file']
		]
		product_selection_df['sector'] = [
			sector_yml['name'] for sector_yml in product_selection_df['sector_yml']
		]

		# Add description
		product_selection_df['description'] = [
			product_yml[content_id]
			for product_yml, content_id in product_selection_df[['product_yml', 'content_id']].to_numpy()
		]

		# Combine type
		product_selection_table = product_selection_df[['sector', 'product', 'type', 'description']].drop_duplicates()
		product_selection_table = _convert_to_bullet_points(product_selection_table, 'type')

		# Create datatable
		product_selection_table = data_loader.rename_user_selection_data_columns(product_selection_table)
		product_selection_table = data_loader.create_data_table(product_selection_table, ['Type'], ['Description'])

		product_layout = html.Div([
			html.P('The following rows contribute:'),
			html.Div(product_selection_table)
		])
	return product_layout

def remove_parents_with_empty_children(report_content):
	if not isinstance(report_content, (list, tuple)):
		return report_content

	updated_children = []
	for node in report_content:
		if hasattr(node, 'children'):
			new_children = remove_parents_with_empty_children(node.children)
			if new_children:
				updated_node = node.__class__(new_children, **{k: v for k, v in node.__dict__.items() if k in ['className', 'style', 'id']})
				updated_children.append(updated_node)
		else:
			updated_children.append(node)
	updated_children = [c for c in updated_children if c is not None and c != [] and c != '']
	if updated_children:
		if hasattr(updated_children[0], 'children'):
			if (len(updated_children) == 1 and hasattr(updated_children[0], '__class__')
					and updated_children[0].__class__.__name__ in {'H1', 'H2', 'H3', 'H4', 'H5', 'H6'}):
				updated_children = []
	return updated_children

def clean_up_sector_overview_and_detail(report_content):
	# Get sector overview and detail layout and its index
	sector_overview_mapping = {i: x for i, x in enumerate(report_content) if x.children[0].children == 'Sector Overview'}
	sector_overview_index = next(iter(sector_overview_mapping.keys()))
	sector_overview_layout = next(iter(sector_overview_mapping.values()))
	sector_detail_mapping = {i: x for i, x in enumerate(report_content) if x.children[0].children == 'Sector Detail'}
	sector_detail_index = next(iter(sector_detail_mapping.keys()))
	sector_detail_layout = next(iter(sector_detail_mapping.values()))

	# Extract product table
	sector_desc = []
	product_desc = sector_overview_layout.children[2].children[0]
	product_datatable = sector_overview_layout.children[2].children[1].children.children[0]
	product_datatable_columns = [col['name'] for col in product_datatable.columns]
	product_selection_df = pd.DataFrame(product_datatable.data, columns=product_datatable_columns)

	# Clean up sector layout
	for sector_section in sector_overview_layout.children[1].children:
		if hasattr(sector_section, 'children'):
			sector_description_layout = sector_section.children
			sector_name = sector_description_layout[0].children

			# Filter product selection for this sector
			filtered_product_selection_df = product_selection_df[product_selection_df['Sector'] == sector_name]
			filtered_product_selection_table = data_loader.create_data_table(
				filtered_product_selection_df, ['Type'], ['Description']
			)
			product_layout = [product_desc] + [filtered_product_selection_table]

			# Get sector-scenario for this sector
			sector_scenario_layout = [
				x for x in sector_detail_layout.children[1].children if x.children[0].children == sector_name
			][0].children[1:]
			sector_scenario_layout = [
				y for x in sector_scenario_layout for y in [html.H3(x.children[0].children), *x.children[1:]]
			]

			# Define sector layout
			sector_layout = html.Div([
				html.H2(sector_name),
				*sector_description_layout[1:],
				*product_layout,
				*sector_scenario_layout,
			])
			sector_desc.append(sector_layout)

	all_sectors_layout = html.Div([
		html.H1('Sector Detail'),
		*sector_desc,
	])
	report_content[sector_overview_index] = all_sectors_layout
	report_content.pop(sector_detail_index)
	return report_content

def create_sidebar_layout(report_content):
	report_content, nav_groups = _extract_navigation_groups(report_content)
	accordion_items = [
		dbc.AccordionItem(
			dbc.Nav(group['links'], vertical=True, pills=True, className="ms-2"),
			title=group['title'],
			item_id=group['id']
		)
		for group in nav_groups
	]
	sidebar_layout = [
		html.H4("Table of Contents", className='text-center p-0 mt-3'),
		dbc.Accordion(accordion_items, start_collapsed=True, always_open=True, flush=True, id="toc-accordion")
	]
	return report_content, sidebar_layout

def _extract_navigation_groups(report_content, current_h1=None, groups=None):
	if groups is None:
		groups = []
	if not isinstance(report_content, (list, tuple)):
		return report_content, groups

	updated_children = []
	for node in report_content:
		if hasattr(node, 'children'):
			child_content = node.children

			if node.__class__.__name__ in {'H1', 'H2'}:
				base_id = str(node.children).replace(' ', '')
				if node.__class__.__name__ == 'H1':
					full_id = base_id
					current_h1 = full_id
					groups.append({'id': full_id, 'title': str(node.children), 'links': []})
					groups[-1]['links'].append(
						dbc.NavLink(str(node.children), href=f"#{full_id}", external_link=True, className="level-1 fw-bold")
					)
				else:
					if current_h1:
						full_id = f"{current_h1}-{base_id}"
					else:
						full_id = base_id
						current_h1 = full_id
						groups.append({'id': full_id, 'title': str(node.children), 'links': []})
						groups[-1]['links'].append(
							dbc.NavLink(str(node.children), href=f"#{full_id}", external_link=True, className="level-1")
						)
					for g in groups:
						if g['id'] == current_h1:
							g['links'].append(
								dbc.NavLink(str(node.children), href=f"#{full_id}", external_link=True, className="level-2")
							)
							break
				params = {k: v for k, v in node.__dict__.items() if k in ['className', 'style']}
				style = (params.get('style') or {}) | {'scrollMarginTop': '120px'}
				params['style'] = style
				updated_node = node.__class__(child_content, id=full_id, **params)
				updated_children.append(updated_node)
			else:
				new_children, groups = _extract_navigation_groups(
					child_content if isinstance(child_content, (list, tuple)) else [child_content],
					current_h1, groups
				)
				if not isinstance(new_children, (list, tuple)):
					new_children = [new_children]
				params = {k: v for k, v in node.__dict__.items() if k in ['className', 'style', 'id']}
				updated_children.append(node.__class__(new_children, **params))
		else:
			updated_children.append(node)

	return updated_children, groups

def _filter_yml_by_scenario(yml, scenario, scenario_mapping_df):
	filtered_scenario_mapping_df = scenario_mapping_df[scenario_mapping_df['scenario_name'] == scenario]
	scenario_risk_type = filtered_scenario_mapping_df['risk_type'].iloc[0]
	scenario_risk_level = filtered_scenario_mapping_df['risk_level'].iloc[0]
	return yml[scenario_risk_type][scenario_risk_level]

def _filter_user_selection_by_section(user_selection_df, section, sub_section):
	filtered_user_selection_df = user_selection_df[
		(user_selection_df['output_structure'] == section)
		& (user_selection_df['sub_section'] == sub_section)
	]
	return filtered_user_selection_df

def _convert_to_bullet_points(df, bullet_point_column_name):
	groupby_variables_list = [x for x in df if x != bullet_point_column_name]
	output_df = df.sort_values(bullet_point_column_name)
	output_df[bullet_point_column_name] = output_df[bullet_point_column_name].astype(str)
	output_df = output_df.groupby(
		groupby_variables_list, as_index=False).agg({bullet_point_column_name: '\n- '.join})
	output_df[bullet_point_column_name] = '- ' + output_df[bullet_point_column_name]
	return output_df

def _check_headers_in_markdown(node, headers_list):
	return len([x.lower() for x in headers_list if x.lower() in node.__dict__['className']]) > 0

def _dash_html_to_string(output_structure_layout):
	# Recursively convert Dash HTML components to HTML string
	if isinstance(output_structure_layout, list):
		return ''.join([_dash_html_to_string(child) for child in output_structure_layout])
	if hasattr(output_structure_layout, 'children'):
		tag = output_structure_layout.__class__.__name__.lower()
		children = _dash_html_to_string(output_structure_layout.children)
		return f"<{tag}>{children}</{tag}>"
	return str(output_structure_layout)

# def _dash_html_to_string(dash_component):
# 	# Recursively convert Dash HTML components to HTML string
# 	if isinstance(dash_component, list):
# 		return ''.join([_dash_html_to_string(child) for child in dash_component])
# 	if hasattr(dash_component, 'props'):
# 		tag = dash_component.__class__.__name__.lower()
# 		children = _dash_html_to_string(dash_component.props.get('children', ''))
# 		return f"<{tag}>{children}</{tag}>"
# 	return str(dash_component)
#
# def _html_to_docx(html_string):
# 	doc = Document()
# 	soup = BeautifulSoup(html_string, "html.parser")
# 	for elem in soup.find_all(['h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'p', 'div', 'ul', 'ol', 'li']):
# 		text = elem.get_text()
# 		if elem.name.startswith('h'):
# 			level = int(elem.name[1])
# 			doc.add_heading(text, level=level)
# 		elif elem.name in ['ul', 'ol']:
# 			for li in elem.find_all('li'):
# 				doc.add_paragraph(li.get_text(), style='List Bullet' if elem.name == 'ul' else 'List Number')
# 		elif elem.name == 'p':
# 			doc.add_paragraph(text)
# 		elif elem.name == 'div':
# 			doc.add_paragraph(text)
# 	return doc
